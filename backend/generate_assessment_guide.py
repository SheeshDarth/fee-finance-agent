from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = REPO_ROOT / "FeeOps_Agent_Demonstration_and_Explanation_Guide.docx"

NAVY = "123047"
BLUE = "2E74B5"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF7E0"
MUTED = "5B6870"
WHITE = "FFFFFF"


def set_font(run, size: float, color: str = NAVY, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 140, bottom: int = 100, end: int = 140) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    width = cell_properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        cell_properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_properties = table._tbl.tblPr
    layout = table_properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    table_width = table_properties.first_child_found_in("w:tblW")
    table_width.set(qn("w:w"), str(sum(widths_dxa)))
    table_width.set(qn("w:type"), "dxa")

    table_indent = table_properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:w"), "140")
    table_indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for grid_column in list(grid):
        grid.remove(grid_column)
    for width_dxa in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width_dxa))
        grid.append(grid_column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_text(paragraph, text: str, *, bold: bool = False, color: str = NAVY, size: float = 10.5, italic: bool = False) -> None:
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)


def add_body(doc: Document, text: str = ""):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if text:
        add_text(paragraph, text)
    return paragraph


def add_labeled(doc: Document, label: str, text: str) -> None:
    paragraph = add_body(doc)
    add_text(paragraph, f"{label}: ", bold=True, color=BLUE)
    add_text(paragraph, text)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style="Heading 1" if level == 1 else "Heading 2")
    paragraph.add_run(text)


def add_callout(doc: Document, title: str, body: str, fill: str = PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title, bold=True, color=NAVY, size=10.5)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.2
    add_text(paragraph, body, color=NAVY, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_show_say_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    headers = ("Show", "Say", "Why it matters")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, NAVY)
        paragraph = cell.paragraphs[0]
        add_text(paragraph, header, bold=True, color=WHITE, size=10)
    for show, say, reason in rows:
        cells = table.add_row().cells
        for index, value in enumerate((show, say, reason)):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            add_text(paragraph, value, size=9.25, color=NAVY, bold=index == 0)
    set_table_geometry(table, [2100, 3300, 3960])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def set_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)

    heading_one = doc.styles["Heading 1"]
    heading_one.base_style = normal
    heading_one.font.name = "Calibri"
    heading_one._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_one._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_one.font.size = Pt(16)
    heading_one.font.bold = True
    heading_one.font.color.rgb = RGBColor.from_string(BLUE)
    heading_one.paragraph_format.space_before = Pt(16)
    heading_one.paragraph_format.space_after = Pt(7)
    heading_one.paragraph_format.keep_with_next = True

    heading_two = doc.styles["Heading 2"]
    heading_two.base_style = normal
    heading_two.font.name = "Calibri"
    heading_two._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_two._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_two.font.size = Pt(12.5)
    heading_two.font.bold = True
    heading_two.font.color.rgb = RGBColor.from_string(NAVY)
    heading_two.paragraph_format.space_before = Pt(11)
    heading_two.paragraph_format.space_after = Pt(5)
    heading_two.paragraph_format.keep_with_next = True

    if "Guide Caption" not in [style.name for style in doc.styles]:
        style = doc.styles.add_style("Guide Caption", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style.font.size = Pt(9)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(header, "FEEOPS | ASSESSMENT DEMONSTRATION GUIDE", bold=True, color=MUTED, size=8.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer, "Scope: assessment prototype. Monetary decisions remain deterministic and reviewable.", color=MUTED, size=8.5)


def build_document() -> None:
    doc = Document()
    set_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(3)
    add_text(title, "FeeOps Agent", bold=True, color=NAVY, size=28)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    add_text(subtitle, "Demonstration & Explanation Guide", bold=True, color=BLUE, size=17)

    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_after = Pt(18)
    add_text(metadata, "Assessment walkthrough | Updated 15 August 2026 | Recommended core demo: 6 minutes", color=MUTED, size=10)

    add_callout(
        doc,
        "The one-line positioning",
        "FeeOps is an auditable school-fee workflow: deterministic Python calculates the ledger and flags uncertainty; Gemini can word a reminder but cannot create, change, approve, or send a monetary decision.",
    )

    add_heading(doc, "1. Before You Begin")
    add_labeled(doc, "Primary goal", "Demonstrate engineering judgement: explain how the workflow protects financial truth before showing the AI layer.")
    add_labeled(doc, "Local preparation", "Run the deterministic backend snapshot, then start the React frontend. The local experience is the most reliable assessment path.")
    add_labeled(doc, "Live preparation", "Only use the Firestore flow when the signed-in reviewer has a reviewers/{uid} document with active: true and the separate Python Firestore worker is running.")
    add_labeled(doc, "Do not share", "Do not put Firebase passwords, API keys, or service-account paths on a slide, in a document, or in a Git commit.")

    code = doc.add_paragraph()
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(10)
    code.paragraph_format.left_indent = Inches(0.15)
    add_text(code, ".\\backend\\venv\\Scripts\\python.exe backend\\agent_runner.py --as-of 2026-08-13\ncd frontend\nnpm run dev", color=NAVY, size=9.5)

    add_heading(doc, "2. The Core Six-Minute Walkthrough")
    add_show_say_table(doc, [
        (
            "Opening (30 sec)",
            "Schools need reliable collection operations, not a model that improvises balances. FeeOps separates deterministic finance logic from AI wording.",
            "Sets the standard for every later claim."
        ),
        (
            "Overview", 
            "The verified ledger shows net due Rs. 121,335, collected Rs. 55,500, outstanding Rs. 65,835, and overdue Rs. 47,835. Late fees are policy-derived.",
            "Shows the result of fee heads, instalments, concessions, waivers, and confident payments."
        ),
        (
            "Payment review",
            "P003 is possible and P004 is unmatched. Together Rs. 19,000 stays out of official collections until a reviewer decides.",
            "Demonstrates the confidence gate: uncertain money is visible, never silently posted."
        ),
        (
            "Collection worklist",
            "The worklist combines outstanding value, ageing, payment history, and plan compliance. Kabir is high risk, but contact is suppressed because an approved plan exists.",
            "Shows explainable prioritisation rather than opaque AI ranking."
        ),
        (
            "Reminder drafts",
            "Aarav receives a 31-60 day firm draft; Meera receives a 0-30 day polite draft. Every draft must preserve the ledger amount and due date exactly.",
            "Shows the bounded AI task and review-only communication policy."
        ),
        (
            "Audit trail",
            "The run records 17 timestamped events, including approvals, reconciliation decisions, positions, and reminder drafts.",
            "Closes the loop on accountability and traceability."
        ),
    ])

    add_heading(doc, "3. Explain the Financial Controls")
    add_labeled(doc, "Arithmetic", "All monetary calculations use integer paise. Rupee strings are display values only.")
    add_labeled(doc, "Allocation", "Late fees are calculated from policy and due date. Concessions, waivers, and CONFIDENT payments are applied deterministically in FIFO order.")
    add_labeled(doc, "Reconciliation", "Only CONFIDENT matches affect verified collections. POSSIBLE and NEEDS_REVIEW rows remain outside the ledger and enter a human work queue.")
    add_labeled(doc, "Reminder safety", "Gemini receives deterministic facts for wording only. A validator rejects a message if it changes or invents a currency amount or due date. No reminder is sent by this prototype.")
    add_labeled(doc, "Auditability", "Approving authority is retained for concessions, waivers, payment plans, reviewer actions, and run completion.")

    add_callout(
        doc,
        "Use this sentence when challenged about hallucinations",
        "The model is not a ledger engine. Python calculates the financial truth first. The model may only phrase facts that the validator checks against that truth.",
        PALE_GOLD,
    )

    add_heading(doc, "4. Technical Explanation for an Interviewer")
    add_labeled(doc, "React dashboard", "The frontend presents the deterministic local snapshot by default and can subscribe to a Firestore run and its child collections when an authenticated reviewer starts a live run.")
    add_labeled(doc, "Deterministic backend", "agent_runner.py orchestrates seed data, reconciliation, ledger positions, worklist scoring, reminders, and audit events. ledger.py owns money calculations; reconciliation.py owns match confidence; reminders.py owns review drafts.")
    add_labeled(doc, "Google ADK", "feeops_adk exposes two read-only tools: run_fee_workflow and get_money_guardrail. The agent instruction explicitly forbids inventing, rounding, altering, or claiming to send money-related information.")
    add_labeled(doc, "Cloud Run", "The deployed Cloud Run service hosts the ADK API and is protected with IAM identity tokens. It runs the packaged deterministic workflow; it is not a Firestore queue worker.")
    add_labeled(doc, "Firestore worker", "firestore_worker.py is a separate polling process. It changes PENDING runs to RUNNING and then AWAITING_REVIEW, writes child collections, and records reviewer actions in the audit trail.")

    add_heading(doc, "5. Optional Live Demonstrations")
    add_labeled(doc, "Firestore reviewer flow", "Sign in with a reviewer account only after confirming the reviewer document contains active: true. Create a run, keep the worker running, and show PENDING to RUNNING to AWAITING_REVIEW. Then record a reviewer action and show REVIEW_ACTION_APPLIED in the audit trail.")
    add_labeled(doc, "Cloud Run API", "Use this only as a technical extension. Obtain an IAM identity token, call the deployed ADK endpoint, and explain that its tools are read-only. Do not imply that this is a confirmed Managed Agent Runtime deployment.")
    add_labeled(doc, "Gemini wording", "If Vertex AI is available, show generationSource: GEMINI and validationPassed: true. If it is unavailable, show the deterministic fallback and explain that safety is unchanged.")
    add_labeled(doc, "Excel upload", "Treat this as a synthetic extension, not primary assessment evidence. The current upload path forwards worksheet JSON to Firestore but does not yet enforce a production schema, duplicate-payment control, or financial import validation.")

    add_heading(doc, "6. Questions You Are Likely To Get")
    add_labeled(doc, "Why not let Gemini calculate fees?", "Because this is financial data. Deterministic, testable code is appropriate for amounts; AI adds value in constrained explanation and wording.")
    add_labeled(doc, "Why is P003 not collected?", "Its narration suggests a family match but lacks a reliable invoice or student reference. The system prefers a visible human review over a wrong ledger posting.")
    add_labeled(doc, "How are families ranked?", "The score is explainable: overdue amount, ageing, late history, partial history, missed history, and average delay. An approved plan can suppress contact.")
    add_labeled(doc, "What happens when the model fails?", "The workflow uses a deterministic template and retains the same amount/date validation. Finance output remains complete.")
    add_labeled(doc, "Is this production-ready?", "It is a strong scoped prototype. Production work would add validated bank imports, duplicate and overpayment handling, identity governance, immutable retention, monitoring, and deployment hardening.")

    add_heading(doc, "7. Claims to Avoid")
    add_labeled(doc, "Do not say", "The agent sends parent messages. It creates drafts for accounts-office review only.")
    add_labeled(doc, "Do not say", "The company Firebase project is live, unless Firebase registration and Native Firestore creation in intern-bnmit-july-2026 have been independently verified.")
    add_labeled(doc, "Do not say", "Cloud Run monitors Firestore. The worker is a separate process in the current implementation.")
    add_labeled(doc, "Do not say", "Excel upload is production ingestion. It is currently a synthetic demo extension without import validation.")
    add_labeled(doc, "Do not say", "The ADK agent changes fees or approves payments. Its exposed tools are read-only.")

    add_heading(doc, "8. Final Pre-Demo Checklist")
    add_labeled(doc, "Check 1", "Backend test command passes all seven financial invariants.")
    add_labeled(doc, "Check 2", "Frontend production build succeeds and the local dashboard opens.")
    add_labeled(doc, "Check 3", "The visible snapshot contains P003 and P004 in Payment review, two reminder drafts, and 17 audit events.")
    add_labeled(doc, "Check 4", "For the live path, confirm Firebase configuration, reviewer active: true, worker process, and correct project credentials.")
    add_labeled(doc, "Check 5", "Keep the conversation focused on financial controls, explainability, and deliberate AI boundaries rather than claiming unfinished cloud or import features.")

    ending = doc.add_paragraph()
    ending.paragraph_format.space_before = Pt(14)
    ending.paragraph_format.space_after = Pt(0)
    ending.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(ending, "A reliable finance agent earns trust by knowing what it may automate, what it must validate, and what a human must still decide.", bold=True, color=NAVY, size=11)

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
