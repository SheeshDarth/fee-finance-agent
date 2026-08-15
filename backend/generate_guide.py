import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_demo_guide():
    doc = Document()

    # Title
    title = doc.add_heading('FeeOps Agent: Demonstration & Explanation Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section 1: Project Overview
    doc.add_heading('1. Project Overview & Pitch', level=1)
    p = doc.add_paragraph(
        "Start the demonstration by explaining the core problem: Schools and institutions struggle with fee collection "
        "because reconciling bank transfers with student accounts is tedious, and chasing overdue payments damages relationships. "
    )
    p.add_run("FeeOps solves this by using AI to automate the entire accounts-receivable workflow.").bold = True

    p2 = doc.add_paragraph()
    p2.add_run("The Secret Sauce (Important!): ").bold = True
    p2.add_run("Explain that while most AI agents hallucinate numbers, this agent uses a ")
    p2.add_run("Deterministic Guardrail Pattern").bold = True
    p2.add_run(". The exact amounts, due dates, and maths are calculated purely by Python. The LLM (Gemini) is only allowed to 'read' these deterministic facts to write polite emails. It cannot change the numbers.")

    # Section 2: How to Demo the Frontend
    doc.add_heading('2. Step-by-Step Demo Flow', level=1)

    doc.add_heading('Step 1: The Base Snapshot', level=2)
    doc.add_paragraph("1. Open the React frontend. Point out the 'Dashboard' showing total overdue amounts.")
    doc.add_paragraph("2. Navigate to 'Collection Worklist'. Show how the agent has automatically ranked families by priority (e.g., who to contact first).")
    doc.add_paragraph("3. Show the 'Audit Trail'. Emphasize that every single calculation and action is logged deterministically for the finance team.")

    doc.add_heading('Step 2: AI Reminder Generation', level=2)
    doc.add_paragraph("1. Navigate to the 'Reminders' tab.")
    doc.add_paragraph("2. Show the drafts (like the one for Aarav Sharma).")
    doc.add_paragraph("3. Point out the green badge that says: 'Gemini wording passed deterministic amount and due-date validation.'")
    doc.add_paragraph("4. Explain: 'The AI drafted a personalized, polite email based on the student's history, but before showing it to the user, our backend validated that the AI didn't hallucinate the Rs. 13,700 amount.'")

    doc.add_heading('Step 3: The Custom Excel Upload (The "Wow" Moment)', level=2)
    doc.add_paragraph("1. Click the 'Download Template' button.")
    doc.add_paragraph("2. Open the 'synthetic_demo.xlsx' file (or the template) in front of them.")
    doc.add_paragraph("3. Show them that you are modifying the data live (e.g., changing a student name or adding a massive payment).")
    doc.add_paragraph("4. Click 'Upload Excel' and upload the file.")
    doc.add_paragraph("5. Wait a few seconds and watch the UI update automatically! The dashboard numbers will change, and the audit trail will reflect the new file.")

    # Section 3: Technical Architecture
    doc.add_heading('3. Technical Architecture (For Technical Judges)', level=1)
    doc.add_paragraph("- Frontend: React (Vite) with real-time Firebase Firestore subscriptions.")
    doc.add_paragraph("- Backend API: Google Agent Development Kit (ADK) running on Cloud Run.")
    doc.add_paragraph("- Background Worker: A Python script monitoring Firestore, generating ledgers, and validating AI outputs.")
    doc.add_paragraph("- AI Model: Gemini 2.5 Flash used exclusively for natural language generation (NLG), sandboxed from core arithmetic.")

    # Save the document
    out_path = Path(r"C:\Users\Siddharth\Desktop\Subhanu\FDE AGENT\FeeOps_Demo_Guide.docx")
    doc.save(out_path)
    print(f"Guide successfully generated at: {out_path}")

if __name__ == "__main__":
    create_demo_guide()
